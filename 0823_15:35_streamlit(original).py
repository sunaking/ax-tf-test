# streamlit_front.py
import streamlit as st
import requests
import uuid
import os
import io
import re

# =============== 기본 설정 ===============
st.set_page_config(page_title="애저딱칼센", layout="wide")

# =============== 공통 스타일 ===============
st.markdown("""
<style>
:root{
  --brand:#00B8CC; --brand-ink:#008080; --ink:#e5e7eb; --muted:#9ca3af;
  --bg:#0b0f14; --card:#0f1620; --mark:#ffe58a; --border:#1f2937;
}
html, body, .stApp{ background:var(--bg) !important; color:var(--ink) !important; }

/* 헤더/사이드바 톤 */
header[data-testid="stHeader"] { background-color: #00B8CC !important; height: 70px; }
[data-testid="stSidebar"] { background-color: #EBFCFF !important; color: black !important; }
[data-testid="stSidebar"] * { color: black !important; }

/* 문서 타이포그래피 */
h1,h2,h3,h4 { line-height:1.25; margin:.6em 0 .4em; font-weight:800; }
h1 { color:var(--brand); font-size: 2.2rem; }
h2 { color:var(--brand-ink); font-size: 1.6rem; }
h3 { color:var(--ink); font-size: 1.25rem; }
p, li { font-size: 1.05rem; }
ul, ol{ padding-left: 1.15rem; } ul li{ margin: 6px 0; } ul li::marker{ color: var(--brand-ink); }
a, a:visited{ color: var(--brand); text-decoration: none; border-bottom: 1px dashed transparent; }
a:hover{ border-bottom-color: var(--brand); }
hr{ border:none; height:1px; background:var(--border); margin: 20px 0; }
blockquote{
  background:linear-gradient(135deg, rgba(0,184,204,.06), rgba(0,128,128,.04));
  border:1px solid var(--border); border-left:6px solid var(--brand);
  padding:12px 14px; border-radius:10px; margin:14px 0;
}
table{ width:100%; border-collapse: collapse; margin: 12px 0; }
th, td{ border: 1px solid var(--border); padding: 10px; }
th{ background:rgba(255,255,255,0.02); color: var(--brand-ink); }
code{ background:rgba(255,255,255,0.04); border:1px solid var(--border); padding:2px 6px; border-radius:6px; }
pre code{ display:block; padding:14px; overflow-x:auto; line-height:1.4; }
mark{ background:var(--mark); color:#222; padding:0 4px; border-radius:4px; }

/* 히어로 */
.hero{
  background:linear-gradient(135deg, rgba(0,184,204,.18), rgba(0,128,128,.10));
  border:1px solid var(--border); border-radius:24px; padding:22px 20px; margin: 6px 0 16px;
}
.hero h2{ margin:0 0 6px 0; color:#5ee7f2; }
.hero p{ margin:0; color:var(--muted); }

/* 툴바(투명) — 검색창 위아래 이상한 박스 제거 */
.toolbar{ padding:0; margin:0; border:none; background:transparent; box-shadow:none; }
.toolbar .stTextInput>div>div{
  background: rgba(255,255,255,0.03);
  border:1px solid var(--border); border-radius:12px;
}
.toolbar .stTextInput input{ height:44px; line-height:44px; }
.toolbar .stButton>button, .toolbar .stDownloadButton>button{
  height:44px; border-radius:12px; padding:0 14px; margin-top:0 !important;
  border:1px solid var(--border); background:none; color:var(--ink);
}

/* st.container(border=True)를 카드처럼 보이게 */
div[data-testid="stContainer"] > div {
  border-radius: 16px !important;
  border: 1px solid var(--border) !important;
  background: linear-gradient(135deg, rgba(0,184,204,.06), rgba(0,128,128,.03)) !important;
  box-shadow: 0 6px 18px rgba(0,0,0,.22) !important;
  padding: 18px !important;
  margin: 14px 0 !important;
}
</style>
""", unsafe_allow_html=True)

# =============== 세션 준비 ===============
if "messages" not in st.session_state:
    st.session_state.messages = []
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# =============== 상수/유틸 ===============
API_SERVER = "http://40.82.143.146/api"   # 필요 시 수정

def highlight_keyword(md_text: str, keyword: str) -> str:
    if not keyword.strip():
        return md_text
    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
    return pattern.sub(lambda m: f"<mark>{m.group(0)}</mark>", md_text)

def md_to_txt_bytes(md_text: str) -> bytes:
    return md_text.encode("utf-8")

def md_to_docx_bytes(md_text: str) -> bytes:
    """간단 규칙 변환: # 헤더→Heading, 구분선→빈 줄, 나머지→문단"""
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx 미설치 또는 임포트 실패") from e
    doc = Document()
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 6)
            content = line[level:].strip() or " "
            doc.add_heading(content, level=level)
        elif re.match(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", line):
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line if line else " ")
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

# =============== 사이드바 ===============
st.sidebar.header("Side bar")
mode = st.sidebar.radio("기능", ["장바구니", "질문하기", "가이드라인"])

# =============== 장바구니 ===============
if mode == "장바구니":
    st.header("🛍️  Azure 상품 장바구니 담기")
    st.write("여기는 성은과장님 개발 예정인 부분이라 아직 내용은 없습니다~")

# =============== 질문하기 ===============
elif mode == "질문하기":
    st.header("📝 Azure 기초 견적 계산 질문")

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"], unsafe_allow_html=True)

    user_input = st.chat_input("질문을 입력하세요:")
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        payload = {"question": user_input, "session_id": st.session_state.session_id}
        with st.spinner("답변 생성중..."):
            try:
                res = requests.post(f"{API_SERVER}/answer", json=payload, timeout=60)
                res.raise_for_status()
                result = res.json()
                answer = result.get("answer", "")
                references = result.get("references", [])

                response = f"**답변:** {answer}"
                st.session_state.messages.append({"role": "assistant", "content": response})

                with st.chat_message("assistant"):
                    st.markdown(response, unsafe_allow_html=True)
                    if references:
                        with st.expander("참고 문서 보기"):
                            for i, ref in enumerate(references, 1):
                                st.markdown(f"{i}. {ref}")
            except requests.exceptions.RequestException as e:
                st.error(f"FastAPI 응답 실패: {e}")

    if len(st.session_state.get("messages", [])) > 1:
        st.markdown("""
            <div style="display: flex; align-items: center; text-align: center; margin: 20px 0;">
                <hr style="flex-grow: 1; border: 0.5px solid var(--border); margin: 0 10px;">
                <span style="white-space: nowrap; color: var(--muted); font-size: 14px;">이전 대화</span>
                <hr style="flex-grow: 1; border: 0.5px solid var(--border); margin: 0 10px;">
            </div>
        """, unsafe_allow_html=True)

# =============== 가이드라인 ===============
elif mode == "가이드라인":
    st.header("📚 KT 가이드라인")

    guide_path = os.path.join(os.getcwd(), "kt_guide.md")
    if not os.path.exists(guide_path):
        st.warning("`kt_guide.md` 파일을 현재 디렉터리에 넣어주세요.")
    else:
        with open(guide_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        # 히어로
        st.markdown("""
        <div class="hero">
          <h2>KT 가이드 문서</h2>
          <p>키워드 검색, 하이라이트, 다양한 포맷으로 내보내기를 지원합니다.</p>
        </div>
        """, unsafe_allow_html=True)

        # 툴바(투명) - 중앙 정렬
        try:
            c1, c2, c3, c4 = st.columns([3, 1, 1, 1], vertical_alignment="center")
        except TypeError:
            c1, c2, c3, c4 = st.columns([3, 1, 1, 1])

        st.markdown('<div class="toolbar">', unsafe_allow_html=True)
        with c1:
            keyword = st.text_input(
                "🔎 키워드 검색 (대소문자 무시)",
                value="",
                placeholder="예: VM, Storage, 가격",
                label_visibility="collapsed",
            )
        with c2:
            st.download_button("TXT로 내보내기",
                data=md_to_txt_bytes(md_text),
                file_name="kt_guide.txt",
                mime="text/plain",
                use_container_width=True
            )
        with c3:
            make_docx = st.button("DOCX로 내보내기", use_container_width=True)
        with c4:
            st.download_button("MD 원문 다운로드",
                data=md_text.encode("utf-8"),
                file_name="kt_guide.md",
                mime="text/markdown",
                use_container_width=True
            )
        st.markdown('</div>', unsafe_allow_html=True)

        # DOCX 변환 버튼 처리
        if make_docx:
            try:
                docx_bytes = md_to_docx_bytes(md_text)
                st.download_button("다운로드: kt_guide.docx",
                    data=docx_bytes,
                    file_name="kt_guide.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception:
                st.error("DOCX 변환 실패: `pip install python-docx` 설치 후 다시 시도하세요.")

        # 검색 결과 요약
        if keyword.strip():
            matches = []
            for idx, line in enumerate(md_text.splitlines(), start=1):
                if re.search(re.escape(keyword), line, re.IGNORECASE):
                    snippet = (line.strip()[:200] + "...") if len(line.strip()) > 200 else line.strip()
                    matches.append((idx, snippet))
            with st.expander(f"검색 결과 {len(matches)}건"):
                for ln, snippet in matches:
                    st.markdown(f"- **{ln}행**: {snippet}")

        # 본문: 섹션별 카드 렌더링 (## 기준) — st.container(border=True)
        sections = re.split(r'(?m)^##\\s+', md_text)

        intro = sections[0].strip()
        if intro:
            with st.container(border=True):
                st.markdown(highlight_keyword(intro, keyword), unsafe_allow_html=True)

        for sec in sections[1:]:
            lines = sec.splitlines()
            if not lines:
                continue
            title = lines[0].strip()
            body  = "\n".join(lines[1:]).strip()
            with st.container(border=True):
                st.markdown(f"## {title}")
                if body:
                    st.markdown(highlight_keyword(body, keyword), unsafe_allow_html=True)

~                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                            
~                                                                                                           
