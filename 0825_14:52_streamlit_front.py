# streamlit_front.py
import os
import io
import re
import uuid
import requests
import pandas as pd
import streamlit as st

# -----------------------------
# 기본 설정
# -----------------------------
st.set_page_config(page_title="애저딱칼센", layout="wide")

# -----------------------------
# 스타일(CSS)
# -----------------------------
st.markdown("""
<style>
:root{
  --brand:#00B8CC; --brand-ink:#008080; --ink:#e5e7eb; --muted:#9ca3af;
  --bg:#0b0f14; --card:#0f1620; --mark:#ffe58a; --border:#1f2937;
}
html, body, .stApp{ background:var(--bg) !important; color:var(--ink) !important; }
header[data-testid="stHeader"] { background-color: #00B8CC !important; height: 70px; }

/* Sidebar: 라이트톤 + 입력 배경색 */
[data-testid="stSidebar"] { background-color: #EBFCFF !important; color: black !important; }
[data-testid="stSidebar"] * { color: black !important; }
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] > div,
[data-testid="stSidebar"] .stNumberInput input,
[data-testid="stSidebar"] .stTextInput input {
  background:#BEFBFF !important; color:#000 !important; border:1px solid #66cbd6 !important;
}
[data-testid="stSidebar"] label p { color:#000 !important; font-weight:600; }

/* 카드 느낌 컨테이너 */
div[data-testid="stContainer"] > div {
  border-radius: 16px !important;
  border: 1px solid var(--border) !important;
  background: linear-gradient(135deg, rgba(0,184,204,.06), rgba(0,128,128,.03)) !important;
  box-shadow: 0 6px 18px rgba(0,0,0,.22) !important;
  padding: 18px !important; margin: 14px 0 !important;
}

/* ===== Right panel: 수량/계약기간 입력 ===== */
.right-panel .accent-input { margin: 8px 0 12px 0 !important; }
.right-panel .accent-input .stNumberInput input{
  background:#E6F4F1 !important;
  border:1px solid #66cbd6 !important;
  color:#000 !important;
  height:44px !important;
  border-radius:12px !important;
  padding: 0 12px !important;
}
.right-panel .accent-input label { margin-bottom:6px !important; }
.right-panel .accent-input .stNumberInput button { display:none !important; }
.right-panel .accent-input input[type=number]::-webkit-outer-spin-button,
.right-panel .accent-input input[type=number]::-webkit-inner-spin-button{ -webkit-appearance:none; margin:0; }
.right-panel .accent-input input[type=number]{ -moz-appearance:textfield; }

/* ===== 공통 액션 버튼(장바구니/내보내기) ===== */
.action-area .stButton > button,
.export-area .stDownloadButton > button {
  background:#95B0B5 !important;
  color:#000 !important;
  border:none !important;
  border-radius:12px !important;
  height:46px !important;
  font-weight:700 !important;
}
.export-area { padding-top: 8px; }

/* 표 가독성 */
table{ width:100%; border-collapse: collapse; }
th, td{ border: 1px solid var(--border); padding: 10px; }
th{ background:rgba(255,255,255,0.04); color: var(--brand-ink); }
</style>
""", unsafe_allow_html=True)

# -----------------------------
# 세션 상태
# -----------------------------
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if "messages" not in st.session_state:
    st.session_state.messages = []
if "cart" not in st.session_state:
    st.session_state.cart = pd.DataFrame(columns=[
        "category","name","region","attributes","reservation_term",
        "unit_price_month_usd","qty","months",
        "ext_month_krw","ext_month_krw_vat","ext_total_krw","ext_total_krw_vat",
        "memo","tags"
    ])
if "_qty" not in st.session_state:
    st.session_state._qty = 1
if "_aoai_in" not in st.session_state:
    st.session_state._aoai_in = 0.0
if "_aoai_out" not in st.session_state:
    st.session_state._aoai_out = 0.0

# -----------------------------
# API helpers
# -----------------------------
API_SERVER = os.environ.get("PRICER_API", "http://40.82.143.146/api")

@st.cache_data(show_spinner=False)
def api_categories():
    return requests.get(f"{API_SERVER}/categories", timeout=8).json()

@st.cache_data(show_spinner=False)
def api_regions():
    return requests.get(f"{API_SERVER}/regions", timeout=8).json()

@st.cache_data(show_spinner=False)
def api_options(category: str, region: str | None):
    params = {"category": category}
    if region:
        params["region"] = region
    return requests.get(f"{API_SERVER}/options", params=params, timeout=12).json()

@st.cache_data(show_spinner=False)
def api_items(category: str, params: dict) -> pd.DataFrame:
    q = {"category": category}
    q.update({k: v for k, v in params.items() if v not in (None, "")})
    r = requests.get(f"{API_SERVER}/items", params=q, timeout=20)
    r.raise_for_status()
    return pd.DataFrame(r.json())

@st.cache_data(show_spinner=False)
def api_fx(base="USD", target="KRW") -> float:
    return float(requests.get(f"{API_SERVER}/fx", params={"base": base, "target": target}, timeout=8).json()["rate"])

@st.cache_data(show_spinner=False)
def api_config() -> dict:
    try:
        return requests.get(f"{API_SERVER}/config", timeout=5).json()
    except Exception:
        return {"vat_rate": 0.10}

# -----------------------------
# 옵션 유효값 계산 (의존형 드롭다운)
# -----------------------------
def _vm_valids(df: pd.DataFrame) -> dict:
    if df.empty: return {"OS":[], "Cores":[], "MemoryGB":[], "ReservationTerm":[]}
    attrs = df["attributes"].apply(lambda x: x if isinstance(x, dict) else {})
    return {
        "OS": sorted({a.get("OS") for a in attrs if a.get("OS")}),
        "Cores": sorted({int(a.get("Cores")) for a in attrs if a.get("Cores") is not None}),
        "MemoryGB": sorted({int(a.get("MemoryGB")) for a in attrs if a.get("MemoryGB") is not None}),
        "ReservationTerm": sorted({t for t in df["reservation_term"].fillna("").astype(str).tolist() if t.strip()}),
    }

def _disk_valids(df: pd.DataFrame) -> dict:
    if df.empty: return {"DiskCategory":[], "Tier":[], "Provisioned_GiB":[], "Provisioned_IOPS":[], "ReservationTerm":[]}
    attrs = df["attributes"].apply(lambda x: x if isinstance(x, dict) else {})
    return {
        "DiskCategory": sorted({a.get("DiskCategory") for a in attrs if a.get("DiskCategory")}),
        "Tier": sorted({a.get("Tier") for a in attrs if a.get("Tier")}),
        "Provisioned_GiB": sorted({int(a.get("Provisioned_GiB")) for a in attrs if a.get("Provisioned_GiB") is not None}),
        "Provisioned_IOPS": sorted({int(a.get("Provisioned_IOPS")) for a in attrs if a.get("Provisioned_IOPS") is not None}),
        "ReservationTerm": sorted({t for t in df["reservation_term"].fillna("").astype(str).tolist() if t.strip()}),
    }

def _db_valids(df: pd.DataFrame) -> dict:
    if df.empty: return {"serviceName":[], "skuName":[], "ReservationTerm":[]}
    attrs = df["attributes"].apply(lambda x: x if isinstance(x, dict) else {})
    return {
        "serviceName": sorted({a.get("serviceName") for a in attrs if a.get("serviceName")}),
        "skuName": sorted({a.get("skuName") for a in attrs if a.get("skuName")}),
        "ReservationTerm": sorted({t for t in df["reservation_term"].fillna("").astype(str).tolist() if t.strip()}),
    }

# -----------------------------
# 유틸(가이드라인 변환)
# -----------------------------
def md_to_txt_bytes(md_text: str) -> bytes:
    return md_text.encode("utf-8")

def md_to_docx_bytes(md_text: str) -> bytes:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx 미설치 또는 임포트 실패") from e
    doc = Document()
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 6)
            content = line[level:].strip() or " "
            doc.add_heading(content, level=level)
        elif re.match(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", line):
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line if line else " ")
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

# -----------------------------
# 사이드바
# -----------------------------
st.sidebar.header("Side bar")
mode = st.sidebar.radio("기능", ["장바구니", "질문하기", "가이드라인"])

# -----------------------------
# 장바구니 모드
# -----------------------------
if mode == "장바구니":
    st.header("🛍️  Azure 상품 장바구니 담기")

    # ---- 사이드바: 카테고리/리전/옵션(의존형) ----
    with st.sidebar:
        try:
            cats = api_categories()
        except Exception as e:
            st.error(f"카테고리 조회 실패: {e}")
            cats = ["VM", "Disk", "DB", "Databricks", "AOAI"]

        try:
            regions = api_regions()
        except Exception as e:
            st.error(f"리전 조회 실패: {e}")
            regions = ["koreacentral"]

        sel_cat = st.selectbox("대분류(Category)", cats, index=0)
        default_idx = regions.index("koreacentral") if "koreacentral" in regions else 0
        sel_region = st.selectbox("Region", regions, index=default_idx)

        fx = api_fx("USD", "KRW")
        cfg = api_config()
        vat = float(cfg.get("vat_rate", 0.10))
        st.caption(f"오늘의 환율(USD→KRW): {fx:,.2f} · 부가세율: {vat*100:.0f}%")

        params: dict = {"region": sel_region}

        if sel_cat == "VM":
            base_df = api_items("VM", {"region": sel_region})
            valids = _vm_valids(base_df)
            p_os = st.selectbox("OS", [""] + valids["OS"]) or None
            if p_os: params["OS"] = p_os; base_df = api_items("VM", params)
            valids = _vm_valids(base_df)
            p_core = st.selectbox("vCPU", [None] + valids["Cores"])
            if p_core is not None: params["Cores"] = p_core; base_df = api_items("VM", params)
            valids = _vm_valids(base_df)
            p_mem = st.selectbox("MemoryGB", [None] + valids["MemoryGB"])
            if p_mem is not None: params["MemoryGB"] = p_mem; base_df = api_items("VM", params)
            valids = _vm_valids(base_df)
            p_term = st.selectbox("ReservationTerm", [""] + valids["ReservationTerm"]) or None
            if p_term: params["VM_ReservationTerm"] = p_term

        elif sel_cat == "Disk":
            base_df = api_items("Disk", {"region": sel_region})
            valids = _disk_valids(base_df)
            p_cat  = st.selectbox("DiskCategory", [""] + valids["DiskCategory"]) or None
            if p_cat: params["DiskCategory"] = p_cat; base_df = api_items("Disk", params)
            valids = _disk_valids(base_df)
            p_tier = st.selectbox("Tier", [""] + valids["Tier"]) or None
            if p_tier: params["Tier"] = p_tier; base_df = api_items("Disk", params)
            valids = _disk_valids(base_df)
            p_gib  = st.selectbox("Provisioned_GiB", [None] + valids["Provisioned_GiB"])
            if p_gib is not None: params["Provisioned_GiB"] = p_gib; base_df = api_items("Disk", params)
            valids = _disk_valids(base_df)
            p_iops = st.selectbox("Provisioned_IOPS", [None] + valids["Provisioned_IOPS"])
            if p_iops is not None: params["Provisioned_IOPS"] = p_iops; base_df = api_items("Disk", params)
            valids = _disk_valids(base_df)
            p_term = st.selectbox("ReservationTerm", [""] + valids["ReservationTerm"]) or None
            if p_term: params["Disk_ReservationTerm"] = p_term

        elif sel_cat == "DB":
            base_df = api_items("DB", {"region": sel_region})
            valids = _db_valids(base_df)
            p_svc = st.selectbox("serviceName", [""] + valids["serviceName"]) or None
            if p_svc: params["serviceName"] = p_svc; base_df = api_items("DB", params)
            valids = _db_valids(base_df)
            p_sku = st.selectbox("skuName", [""] + valids["skuName"]) or None
            if p_sku: params["skuName"] = p_sku; base_df = api_items("DB", params)
            valids = _db_valids(base_df)
            p_term = st.selectbox("reservationTerm", [""] + valids["ReservationTerm"]) or None
            if p_term: params["DB_ReservationTerm"] = p_term

        elif sel_cat == "Databricks":
            opts = api_options("Databricks", sel_region)
            p_mn = st.selectbox("metername(상품종류)", [""] + opts.get("metername", [])) or None
            if p_mn: params["metername"] = p_mn

        elif sel_cat == "AOAI":
            opts = api_options("AOAI", sel_region)
            p_mn = st.selectbox("metername(상품종류)", [""] + opts.get("metername", [])) or None
            if p_mn: params["metername"] = p_mn

    # ---- 결과 조회 ----
    try:
        df = api_items(sel_cat, params | {"category": sel_cat})
    except Exception as e:
        st.error(f"아이템 조회 실패: {e}")
        df = pd.DataFrame()

    if df.empty:
        st.info("먼저 metername(상품종류)을 선택해 주세요." if sel_cat in ("AOAI","Databricks") else "조건에 맞는 항목이 없습니다. 좌측 필터를 조정해 보세요.")
        st.stop()

    # 라벨 (백엔드가 label 넣지만 방어)
    if "label" not in df.columns:
        def row_label(r):
            parts = [r.get("name",""), r.get("category",""), r.get("region","")]
            attrs = r.get("attributes", {}) if isinstance(r.get("attributes"), dict) else {}
            if attrs.get("OS"): parts.append(str(attrs["OS"]))
            if attrs.get("Cores") is not None: parts.append(f"{int(attrs['Cores'])}vCPU")
            if attrs.get("MemoryGB") is not None: parts.append(f"{int(attrs['MemoryGB'])}GB")
            if r.get("reservation_term"): parts.append(str(r["reservation_term"]))
            return " | ".join([str(x) for x in parts if str(x).strip()])
        df["label"] = df.apply(row_label, axis=1)

    left, right = st.columns([1,1], gap="large")

    with left:
        st.subheader("항목 선택")
        option = st.selectbox("옵션", df["label"].tolist())
        sel = df[df["label"] == option].iloc[0]
        st.write(sel)

    with right:
        st.markdown('<div class="right-panel">', unsafe_allow_html=True)
        st.subheader("수량/기간/금액")

        # AOAI는 입력/출력 토큰을 개별 입력
        aoai_mode = (sel_cat == "AOAI")
        dbr_mode  = (sel_cat == "Databricks")

        if aoai_mode:
            st.markdown('<div class="accent-input">', unsafe_allow_html=True)
            aoai_in = st.number_input("입력 토큰 수 (단위: 토큰)", min_value=0.0, value=float(st.session_state._aoai_in), step=1000.0, format="%.0f")
            st.session_state._aoai_in = aoai_in
            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('<div class="accent-input">', unsafe_allow_html=True)
            aoai_out = st.number_input("출력 토큰 수 (단위: 토큰)", min_value=0.0, value=float(st.session_state._aoai_out), step=1000.0, format="%.0f")
            st.session_state._aoai_out = aoai_out
            st.markdown('</div>', unsafe_allow_html=True)

            # AOAI는 qty 의미가 약하므로 숨김 대신 1로 고정
            qty = 1
        else:
            st.markdown('<div class="accent-input">', unsafe_allow_html=True)
            qty = st.number_input("수량", min_value=1, value=int(st.session_state.get("_qty", 1)),
                                  step=1, key="_qty", label_visibility="visible")
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="accent-input">', unsafe_allow_html=True)
        months = st.number_input("계약기간(개월)", min_value=1, value=12, step=1)
        st.markdown('</div>', unsafe_allow_html=True)

        # 메모 / 태그
        memo = st.text_input("메모(선택)")
        tags = st.text_input("태그(쉼표로 구분, 예: prod,web)")

        # 금액 계산 (KRW만; VAT별도/포함)
        fx = api_fx("USD", "KRW")
        vat = float(api_config().get("vat_rate", 0.10))

        retail_price = float(sel.get("retail_price", 0.0))
        unit_month_usd = float(sel.get("price_per_month_after_rule", sel.get("price_per_month", retail_price)))

        # AOAI/Databricks 특별 산식
        if aoai_mode:
            # unit_of_measure는 보통 '1K tokens'
            # 월 합계(USD) = (retail * in_tokens + retail * out_tokens) / 1000
            unit_month_usd = (retail_price * (st.session_state._aoai_in + st.session_state._aoai_out)) / 1000.0
            ext_month_krw  = unit_month_usd * fx
        elif dbr_mode:
            # 월 환산 730h
            unit_month_usd = retail_price * 730.0
            ext_month_krw  = unit_month_usd * int(qty) * fx
        else:
            # 기본: 월단가(USD) * 수량 * 환율
            ext_month_krw  = unit_month_usd * int(qty) * fx

        ext_total_krw  = ext_month_krw * int(months)
        month_krw_vat = ext_month_krw * (1 + vat)
        total_krw_vat = ext_total_krw * (1 + vat)
        year_krw      = ext_month_krw * 12
        year_krw_vat  = month_krw_vat * 12

        # 요약 표
        summary_df = pd.DataFrame(
            {
                "VAT별도 (KRW)": [round(ext_month_krw), round(year_krw), round(ext_total_krw)],
                "VAT포함 (KRW)": [round(month_krw_vat), round(year_krw_vat), round(total_krw_vat)],
            },
            index=["월 합계", "연간 금액", "총액-계약기간"]
        )
        st.table(summary_df.style.format("{:,.0f}"))

        # 장바구니 담기
        st.markdown('<div class="action-area">', unsafe_allow_html=True)
        if st.button("장바구니 담기", use_container_width=True):
            row = {
                "category": sel.get("category"),
                "name": sel.get("name"),
                "region": sel.get("region"),
                "attributes": sel.get("attributes"),
                "reservation_term": sel.get("reservation_term",""),
                "unit_price_month_usd": round(unit_month_usd, 6),
                "qty": int(qty),
                "months": int(months),
                "ext_month_krw": round(ext_month_krw, 0),
                "ext_total_krw": round(ext_total_krw, 0),
                "ext_month_krw_vat": round(month_krw_vat, 0),
                "ext_total_krw_vat": round(total_krw_vat, 0),
                "memo": memo.strip(),
                "tags": ",".join([t.strip() for t in tags.split(",") if t.strip()])
            }
            st.session_state.cart = pd.concat([st.session_state.cart, pd.DataFrame([row])], ignore_index=True)
            st.toast("장바구니에 담었습니다.", icon="✅")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)  # /right-panel

    # ---- 장바구니 ----
    st.subheader("장바구니")
    if st.session_state.cart.empty:
        st.info("비어 있습니다.")
    else:
        cart = st.session_state.cart.copy()

        def _short_attr(a):
            if not isinstance(a, dict): return ""
            bits=[]
            if a.get("OS"): bits.append(a["OS"])
            if a.get("Cores") is not None: bits.append(f'{int(a["Cores"])}vCPU')
            if a.get("MemoryGB") is not None: bits.append(f'{int(a["MemoryGB"])}GB')
            if a.get("DiskCategory"): bits.append(a["DiskCategory"])
            if a.get("Tier"): bits.append(a["Tier"])
            if a.get("metername"): bits.append(a["metername"])
            return ", ".join(bits)

        cart["attrs"] = cart["attributes"].apply(_short_attr)

        show_cols = [
            "category","name","region","attrs","reservation_term",
            "qty","months",
            "ext_month_krw","ext_month_krw_vat",
            "ext_total_krw","ext_total_krw_vat",
            "memo","tags"
        ]
        st.dataframe(cart[show_cols], use_container_width=True, height=360)

        c1, c2 = st.columns(2)
        c1.metric("월 합계 (KRW, VAT포함)", f"{cart['ext_month_krw_vat'].sum():,.0f}")
        c2.metric("총액(기간, KRW, VAT포함)", f"{cart['ext_total_krw_vat'].sum():,.0f}")

        st.markdown("#### 항목별 액션")
        for i, row in cart.iterrows():
            cols = st.columns([6,2,2])
            with cols[0]:
                st.write(f"**[{i}]** {row['category']} · {row['name']} · {row['region']} · x{row['qty']} ( {row.get('attrs','')} )")
            with cols[2]:
                if st.button("🗑️ 삭제", key=f"del_{i}", use_container_width=True):
                    st.session_state.cart = st.session_state.cart.drop(index=int(i)).reset_index(drop=True)
                    st.rerun()

        with st.expander("내보내기"):
            st.markdown('<div class="export-area">', unsafe_allow_html=True)
            colC, colD = st.columns(2)
            with colC:
                csv_bytes = cart[show_cols].to_csv(index=False).encode("utf-8")
                st.download_button("CSV 다운로드", data=csv_bytes, file_name="cart.csv",
                                   mime="text/csv", use_container_width=True)
            with colD:
                bio = io.BytesIO()
                try:
                    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
                        cart[show_cols].to_excel(writer, index=False, sheet_name="Cart")
                except Exception:
                    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
                        cart[show_cols].to_excel(writer, index=False, sheet_name="Cart")
                bio.seek(0)
                st.download_button("엑셀(xlsx) 다운로드", data=bio.getvalue(),
                    file_name="cart.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

# -----------------------------
# 질문하기 / 가이드라인
# -----------------------------
elif mode == "질문하기":
    st.header("📝 Azure 기초 견적 계산 질문")
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"], unsafe_allow_html=True)

    user_input = st.chat_input("질문을 입력하세요:")
    if user_input:
        st.session_state.messages.append({"role":"user","content":user_input})
        with st.chat_message("user"): st.markdown(user_input)

        payload = {"question": user_input, "session_id": st.session_state.session_id}
        try:
            res = requests.post(f"{API_SERVER}/answer", json=payload, timeout=60)
            res.raise_for_status()
            result = res.json()
            answer = result.get("answer","")
            references = result.get("references",[])
            response = f"**답변:** {answer}"
            st.session_state.messages.append({"role":"assistant","content":response})
            with st.chat_message("assistant"):
                st.markdown(response, unsafe_allow_html=True)
                if references:
                    with st.expander("참고 문서 보기"):
                        for i, ref in enumerate(references, 1):
                            st.markdown(f"{i}. {ref}")
        except requests.exceptions.RequestException as e:
            st.error(f"FastAPI 응답 실패: {e}")

elif mode == "가이드라인":
    st.header("📚 KT 가이드라인")

    guide_path = os.path.join(os.getcwd(), "kt_guide.md")
    if not os.path.exists(guide_path):
        st.warning("`kt_guide.md` 파일을 현재 디렉터리에 넣어주세요.")
    else:
        with open(guide_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        try:
            c1, c2, c3, c4 = st.columns([3,1,1,1], vertical_alignment="center")
        except TypeError:
            c1, c2, c3, c4 = st.columns([3,1,1,1])

        with c1:
            keyword = st.text_input("🔎 키워드 검색 (대소문자 무시)", value="", placeholder="예: VM, Storage, 가격", label_visibility="collapsed")
        with c2:
            st.download_button("TXT로 내보내기", data=md_to_txt_bytes(md_text), file_name="kt_guide.txt", mime="text/plain", use_container_width=True)
        with c3:
            make_docx = st.button("DOCX로 내보내기", use_container_width=True)
        with c4:
            st.download_button("MD 원문 다운로드", data=md_text.encode("utf-8"), file_name="kt_guide.md", mime="text/markdown", use_container_width=True)

        if make_docx:
            try:
                docx_bytes = md_to_docx_bytes(md_text)
                st.download_button("다운로드: kt_guide.docx",
                    data=docx_bytes, file_name="kt_guide.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception:
                st.error("DOCX 변환 실패: `pip install python-docx` 설치 후 다시 시도하세요.")

        if keyword.strip():
            matches = []
            for idx, line in enumerate(md_text.splitlines(), start=1):
                if re.search(re.escape(keyword), line, re.IGNORECASE):
                    snippet = (line.strip()[:200] + "...") if len(line.strip()) > 200 else line.strip()
                    matches.append((idx, snippet))
            with st.expander(f"검색 결과 {len(matches)}건"):
                for ln, snippet in matches:
                    st.markdown(f"- **{ln}행**: {snippet}")

        sections = re.split(r'(?m)^##\\s+', md_text)
        intro = sections[0].strip()
        if intro:
            with st.container(border=True):
                st.markdown(intro, unsafe_allow_html=True)

        for sec in sections[1:]:
            lines = sec.splitlines()
            if not lines: continue
            title = lines[0].strip()
            body  = "\n".join(lines[1:]).strip()
            with st.container(border=True):
                st.markdown(f"## {title}")
                if body:
                    st.markdown(body, unsafe_allow_html=True)
